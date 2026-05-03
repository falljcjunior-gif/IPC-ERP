from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# Styles
def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    if level == 1:
        h.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def para(text, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def colored_cell(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

# ── COVER PAGE ──────────────────────────────────────────────────────
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('RAPPORT D\'AUDIT DE SÉCURITÉ APPLICATIVE')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('IPC ERP — Analyse Statique & Architecture')
r2.font.size = Pt(14)
r2.font.color.rgb = RGBColor(0x55, 0x55, 0x99)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'Date : {datetime.date.today().strftime("%d %B %Y")}\n')
meta.add_run('Auditeur : Antigravity Security Intelligence\n')
meta.add_run('Classification : CONFIDENTIEL')

doc.add_page_break()

# ── 1. RÉSUMÉ EXÉCUTIF ───────────────────────────────────────────────
heading('1. Résumé Exécutif', 1)
para(
    "L'audit de sécurité applicative de l'ERP I.P.C. a été réalisé par analyse statique exhaustive "
    "du code source (React 19, Firebase, Cloud Functions Node.js). Le périmètre couvre l'ensemble "
    "des services frontend, backend (Cloud Functions), règles Firestore, gestion des identités et "
    "dépendances tierces.\n\n"
    "Synthèse des résultats :"
)

summary_table = doc.add_table(rows=5, cols=3)
summary_table.style = 'Table Grid'
headers = ['Criticité', 'Nombre', 'Statut']
colors  = ['1A1A2E', '1A1A2E', '1A1A2E']
for i, h in enumerate(headers):
    cell = summary_table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    colored_cell(cell, 'C0392B') if i == 0 else colored_cell(cell, '2C3E50')
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

data = [
    ('🔴 Critique', '3', 'Action immédiate requise'),
    ('🟠 Haute',    '4', 'Corriger sous 7 jours'),
    ('🟡 Moyenne',  '5', 'Corriger sous 30 jours'),
    ('🟢 Faible',   '3', 'Recommandé'),
]
for row_i, (c, n, s) in enumerate(data, start=1):
    row = summary_table.rows[row_i]
    row.cells[0].text = c
    row.cells[1].text = n
    row.cells[2].text = s

doc.add_paragraph()
para(
    "Avertissement principal : 3 vulnérabilités critiques ont été identifiées nécessitant "
    "une correction immédiate avant tout déploiement en production : "
    "(1) clés Firebase hardcodées dans le Service Worker public, "
    "(2) bypass d'authentification par email fixe dans le code client, "
    "(3) absence de vérification d'authentification sur deleteUserAccount via Custom Claims."
)

doc.add_page_break()

# ── 2. PÉRIMÈTRE & MÉTHODOLOGIE ──────────────────────────────────────
heading('2. Périmètre et Méthodologie', 1)
heading('2.1 Périmètre analysé', 2)
items = [
    'Frontend React 19 — src/ (composants, services, store Zustand, modules ERP)',
    'Backend Firebase — Cloud Functions v2 (nexus, admin, social, triggers, backup)',
    'Sécurité Firestore — firestore.rules',
    'Configuration — .env, firebase/config.js, firebase-messaging-sw.js',
    'Dépendances — package.json (frontend + functions)',
    'Moteur métier — IpcEngine.js, WorkflowEngine.js, firestore.service.js',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

heading('2.2 Méthodologie', 2)
methods = [
    'Analyse statique exhaustive (SAST) — lecture fichier par fichier',
    'Simulation d\'attaques OWASP Top 10 & CWE Top 25',
    'Revue des flux d\'authentification / autorisation / session',
    'Audit des configurations et secrets',
    'Revue des dépendances tierces',
]
for m in methods:
    doc.add_paragraph(m, style='List Bullet')

doc.add_page_break()

# ── 3. TABLEAU RÉCAPITULATIF ──────────────────────────────────────────
heading('3. Tableau Récapitulatif des Vulnérabilités', 1)

cols = ['ID', 'Criticité', 'Fichier', 'Vulnérabilité', 'CWE']
vulns = [
    ('V-01', '🔴 CRITIQUE', 'firebase-messaging-sw.js:4-11', 'Clés Firebase hardcodées dans Service Worker public', 'CWE-798'),
    ('V-02', '🔴 CRITIQUE', 'user.service.js:32 / App.jsx:100', 'Bypass auth par email hardcodé côté client', 'CWE-284'),
    ('V-03', '🔴 CRITIQUE', 'functions/admin.js:34', 'Vérification rôle via Firestore (pas Custom Claims)', 'CWE-269'),
    ('V-04', '🟠 HAUTE',    '.env:19', 'Clé de chiffrement XOR faible dans VITE_STORE_KEY', 'CWE-326'),
    ('V-05', '🟠 HAUTE',    'functions/social.js:101', 'Webhook verify token hardcodé en fallback', 'CWE-798'),
    ('V-06', '🟠 HAUTE',    'nexus.js:51', 'Fallback API key depuis Firestore (contournable)', 'CWE-522'),
    ('V-07', '🟠 HAUTE',    'App.jsx:20-33', 'OAuth code exposé dans URL sans state CSRF', 'CWE-352'),
    ('V-08', '🟡 MOYENNE',  'nexus.js:75-79', 'Injection de prompt possible via erpContext', 'CWE-20'),
    ('V-09', '🟡 MOYENNE',  'triggers.js:242', 'Recherche user par nom (non-unique)', 'CWE-285'),
    ('V-10', '🟡 MOYENNE',  'firestore.rules', 'Collections non couvertes par les règles', 'CWE-284'),
    ('V-11', '🟡 MOYENNE',  'config.js:26', 'experimentalForceLongPolling activé en prod', 'CWE-400'),
    ('V-12', '🟡 MOYENNE',  'store/index.js', 'XOR encryption — pas un vrai chiffrement', 'CWE-327'),
    ('V-13', '🟡 MOYENNE',  'nexus.js:125', 'Message d\'erreur interne exposé au client', 'CWE-209'),
    ('V-14', '🟢 FAIBLE',   'triggers.js:101', 'Audit log sans signature (non répudiable)', 'CWE-778'),
    ('V-15', '🟢 FAIBLE',   'firebase-messaging-sw.js:1', 'Version Firebase SDK 9.0.0 (obsolète)', 'CWE-1035'),
]

table = doc.add_table(rows=len(vulns)+1, cols=len(cols))
table.style = 'Table Grid'

for i, c in enumerate(cols):
    cell = table.rows[0].cells[i]
    cell.text = c
    cell.paragraphs[0].runs[0].bold = True
    colored_cell(cell, '1A1A2E')
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

crit_colors = {'CRITIQUE': 'FADBD8', 'HAUTE': 'FDEBD0', 'MOYENNE': 'FEF9E7', 'FAIBLE': 'EAFAF1'}
for row_i, v in enumerate(vulns, start=1):
    row = table.rows[row_i]
    for ci, val in enumerate(v):
        row.cells[ci].text = val
    key = v[1].split()[-1]
    bg = crit_colors.get(key, 'FFFFFF')
    for ci in range(len(cols)):
        colored_cell(row.cells[ci], bg)

doc.add_page_break()

# ── 4. DÉTAIL DES VULNÉRABILITÉS ─────────────────────────────────────
heading('4. Détail des Vulnérabilités', 1)

findings = [
    {
        'id': 'V-01',
        'crit': '🔴 CRITIQUE',
        'title': 'Clés Firebase Hardcodées dans Service Worker Public',
        'file': 'public/firebase-messaging-sw.js — Lignes 4-11',
        'cwe': 'CWE-798 : Use of Hard-coded Credentials',
        'desc': (
            "Le fichier public/firebase-messaging-sw.js contient en clair la clé API Firebase "
            "(AIzaSyBMQwaE0JnyJ-0zHQI2Ydc2kYD5MiVzoUw), le projectId, le messagingSenderId et l'appId. "
            "Ce fichier est servi publiquement sans authentification par le serveur web."
        ),
        'exploit': (
            "Tout visiteur peut récupérer ce fichier via GET /firebase-messaging-sw.js. "
            "Avec ces clés, un attaquant peut :\n"
            "• Lire les collections Firestore si les règles sont permissives\n"
            "• Abuser du quota Firebase (DoS économique)\n"
            "• Enregistrer des tokens FCM frauduleux\n"
            "• Exploiter les Storage Rules si mal configurées"
        ),
        'fix': (
            "Pour un Service Worker, les variables d'environnement Vite ne sont pas disponibles. "
            "Solution : générer dynamiquement le SW au build via un plugin Vite qui injecte les valeurs, "
            "OU utiliser Firebase App Check (attestation) pour limiter l'usage des clés aux apps légitimes.\n\n"
            "Correctif immédiat : activer Firebase App Check avec reCAPTCHA v3 pour bloquer les appels "
            "non-autorisés même si les clés sont connues."
        ),
    },
    {
        'id': 'V-02',
        'crit': '🔴 CRITIQUE',
        'title': 'Bypass d\'Authentification par Email Hardcodé (Identity Bridge)',
        'file': 'src/services/user.service.js:32, 84 — src/App.jsx:100-107',
        'cwe': 'CWE-284 : Improper Access Control',
        'desc': (
            "La logique d'Identity Bridge octroie SUPER_ADMIN à tout utilisateur Firebase "
            "dont l'email est fall.jcjunior@gmail.com. Cette vérification est effectuée "
            "côté client (navigateur), ce qui la rend contournable."
        ),
        'exploit': (
            "Un attaquant qui compromet le compte Google fall.jcjunior@gmail.com obtient "
            "automatiquement SUPER_ADMIN sans aucun contrôle Firestore. "
            "De plus, la comparaison d'email côté client peut être manipulée via "
            "des outils de debugging (React DevTools, modification mémoire)."
        ),
        'fix': (
            "Remplacer par des Firebase Custom Claims, définis côté serveur uniquement :\n"
            "// Cloud Function (admin.js)\n"
            "await admin.auth().setCustomUserClaims(uid, { role: 'SUPER_ADMIN' });\n\n"
            "// Client — vérification via token ID\n"
            "const token = await user.getIdTokenResult();\n"
            "const role = token.claims.role ?? 'STAFF';\n\n"
            "Supprimer entièrement les checks d'email dans user.service.js et App.jsx."
        ),
    },
    {
        'id': 'V-03',
        'crit': '🔴 CRITIQUE',
        'title': 'Vérification de Rôle Admin via Firestore (Contournable)',
        'file': 'functions/modules/admin.js:31-38',
        'cwe': 'CWE-269 : Improper Privilege Management',
        'desc': (
            "La fonction deleteUserAccount vérifie le rôle SUPER_ADMIN en lisant "
            "directement la collection Firestore users. Un attaquant ayant accès "
            "en écriture à Firestore (ou exploitant une faille de règle) peut "
            "modifier son propre document pour s'auto-élever."
        ),
        'exploit': (
            "Si les règles Firestore permettent à un utilisateur de modifier "
            "son propre champ permissions.roles, il peut s'attribuer SUPER_ADMIN "
            "puis appeler deleteUserAccount sur n'importe quel UID."
        ),
        'fix': (
            "Vérifier via Custom Claims dans le token Firebase Auth, immuables côté client :\n"
            "// ✅ Correct\n"
            "const isSuperAdmin = request.auth.token.role === 'SUPER_ADMIN';\n"
            "if (!isSuperAdmin) throw new HttpsError('permission-denied', '...');"
        ),
    },
    {
        'id': 'V-04',
        'crit': '🟠 HAUTE',
        'title': 'Clé XOR Faible pour Chiffrement du Store Local',
        'file': '.env:19 — src/store/index.js',
        'cwe': 'CWE-326 : Inadequate Encryption Strength',
        'desc': (
            "La clé VITE_STORE_KEY=\"ipc_erp_secure_key_fall_2026\" est statique, "
            "prévisible, et embarquée dans le bundle JavaScript. "
            "L'algorithme XOR n'est pas un chiffrement cryptographique."
        ),
        'fix': (
            "Utiliser l'API SubtleCrypto du navigateur (AES-GCM) avec une clé dérivée "
            "de la session Firebase. Pour les données critiques, ne pas les persister "
            "en localStorage — utiliser sessionStorage ou mémoire uniquement."
        ),
    },
    {
        'id': 'V-05',
        'crit': '🟠 HAUTE',
        'title': 'Webhook Verify Token Hardcodé en Fallback',
        'file': 'functions/modules/social.js:101',
        'cwe': 'CWE-798 : Use of Hard-coded Credentials',
        'desc': (
            "Le token de vérification du webhook Meta a un fallback hardcodé : "
            "'ipc_erp_webhook_secret_2026'. Si META_WEBHOOK_VERIFY_TOKEN n'est "
            "pas défini, n'importe qui connaissant ce secret peut valider un "
            "faux webhook Meta."
        ),
        'fix': (
            "Supprimer le fallback. Lever une exception si la variable n'est pas définie :\n"
            "const TOKEN = process.env.META_WEBHOOK_VERIFY_TOKEN;\n"
            "if (!TOKEN) throw new Error('META_WEBHOOK_VERIFY_TOKEN not configured');"
        ),
    },
    {
        'id': 'V-06',
        'crit': '🟠 HAUTE',
        'title': 'Clé Gemini API Stockée dans Firestore (Fallback)',
        'file': 'functions/modules/nexus.js:51-53',
        'cwe': 'CWE-522 : Insufficiently Protected Credentials',
        'desc': (
            "Si GEMINI_API_KEY n'est pas configuré en Secret Manager, nexus.js "
            "lit la clé depuis Firestore (system_config/ai_config). "
            "Tout utilisateur ayant accès en lecture à cette collection "
            "peut récupérer la clé Gemini."
        ),
        'fix': (
            "Supprimer le fallback Firestore. Utiliser exclusivement Secret Manager. "
            "Ajouter une règle Firestore bloquant tout accès à system_config :\n"
            "match /system_config/{doc} { allow read, write: if false; }"
        ),
    },
    {
        'id': 'V-07',
        'crit': '🟠 HAUTE',
        'title': 'OAuth Callback sans Protection CSRF (State Parameter)',
        'file': 'src/App.jsx:20-33',
        'cwe': 'CWE-352 : Cross-Site Request Forgery',
        'desc': (
            "Le callback OAuth Facebook traite le paramètre 'code' depuis l'URL "
            "sans vérifier un paramètre 'state' aléatoire. "
            "Un attaquant peut forger une URL de callback et voler le code OAuth."
        ),
        'fix': (
            "Générer un state aléatoire avant la redirection OAuth, "
            "le stocker en sessionStorage, et le vérifier au retour :\n"
            "const state = crypto.randomUUID();\n"
            "sessionStorage.setItem('oauth_state', state);\n"
            "// Au callback :\n"
            "if (urlParams.get('state') !== sessionStorage.getItem('oauth_state'))\n"
            "  throw new Error('CSRF détecté');"
        ),
    },
    {
        'id': 'V-08',
        'crit': '🟡 MOYENNE',
        'title': 'Injection de Prompt via erpContext',
        'file': 'functions/modules/nexus.js:75-79',
        'cwe': 'CWE-20 : Improper Input Validation',
        'desc': (
            "Les champs userName, activeModule, kpis sont injectés directement "
            "dans le systemPrompt sans sanitisation. Un utilisateur malveillant "
            "peut injecter des instructions dans le prompt pour détourner Nexus."
        ),
        'fix': (
            "Sanitiser les champs injectés dans le prompt :\n"
            "const safeUserName = userName.replace(/[^a-zA-ZÀ-ÿ0-9 ]/g, '').slice(0, 50);\n"
            "Appliquer des limites strictes via Zod sur la longueur et le format "
            "de chaque champ de erpContext."
        ),
    },
    {
        'id': 'V-09',
        'crit': '🟡 MOYENNE',
        'title': 'Lookup Utilisateur par Nom (Non-Unique)',
        'file': 'functions/modules/triggers.js:242, 286, 373',
        'cwe': 'CWE-285 : Improper Authorization',
        'desc': (
            "Les triggers Butler recherchent les utilisateurs via where('nom', '==', managerName). "
            "Le nom n'étant pas unique, deux employés homonymes recevront les mêmes notifications "
            "et l'un peut intercepter des alertes destinées à l'autre."
        ),
        'fix': (
            "Stocker et utiliser les UIDs (non les noms) pour les références inter-documents. "
            "Remplacer managerName par managerId dans les documents projet."
        ),
    },
    {
        'id': 'V-10',
        'crit': '🟡 MOYENNE',
        'title': 'Collections Firestore Non Couvertes par les Règles',
        'file': 'firestore.rules',
        'cwe': 'CWE-284 : Improper Access Control',
        'desc': (
            "Les nouvelles collections créées dynamiquement (marketing_messages, ai_logs, "
            "system_config, social_tokens, rooms) peuvent ne pas avoir de règles "
            "explicites, héritant du deny-by-default — mais system_config et social_tokens "
            "contiennent des secrets critiques qui méritent une règle explicite."
        ),
        'fix': (
            "Ajouter explicitement :\n"
            "match /system_config/{doc} { allow read, write: if false; }\n"
            "match /social_tokens/{doc} { allow read, write: if false; }\n"
            "match /ai_logs/{doc} { allow read: if isAdmin(); allow write: if false; }"
        ),
    },
    {
        'id': 'V-13',
        'crit': '🟡 MOYENNE',
        'title': 'Message d\'Erreur Interne Exposé au Client',
        'file': 'functions/modules/nexus.js:125',
        'cwe': 'CWE-209 : Information Exposure Through Error Message',
        'desc': (
            "throw new Error(`AI_ERROR: ${error.message}`) expose le message "
            "d'erreur interne Gemini au client, révélant potentiellement "
            "des détails sur l'infrastructure."
        ),
        'fix': (
            "throw new HttpsError('internal', 'Service IA temporairement indisponible.');\n"
            "logger.error('Nexus internal error:', error); // Log côté serveur uniquement"
        ),
    },
    {
        'id': 'V-14',
        'crit': '🟢 FAIBLE',
        'title': 'Audit Logs Sans Signature Cryptographique',
        'file': 'functions/modules/triggers.js:95-115',
        'cwe': 'CWE-778 : Insufficient Logging',
        'desc': (
            "Les audit_logs sont écrits dans Firestore sans signature. "
            "Un SUPER_ADMIN malveillant peut modifier ou supprimer des entrées "
            "d'audit après coup, compromettant la non-répudiation."
        ),
        'fix': (
            "Exporter les logs critiques vers Cloud Logging (immuable) en parallèle. "
            "Optionnellement, hasher chaque entrée avec le contenu du document original."
        ),
    },
    {
        'id': 'V-15',
        'crit': '🟢 FAIBLE',
        'title': 'Version Firebase SDK Obsolète dans Service Worker',
        'file': 'public/firebase-messaging-sw.js:1-2',
        'cwe': 'CWE-1035 : OWASP Top Ten 2017 A9 — Using Components with Known Vulnerabilities',
        'desc': (
            "Le Service Worker importe firebase-app-compat@9.0.0 via CDN gstatic. "
            "La version 9.0.0 est obsolète (version actuelle ~10.x). "
            "Des CVE peuvent exister sur cette version."
        ),
        'fix': "Mettre à jour vers la dernière version stable et épingler la version en production.",
    },
]

for f in findings:
    heading(f"{f['id']} — {f['title']}", 2)
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Criticité'
    table.rows[0].cells[1].text = f['crit']
    table.rows[1].cells[0].text = 'Fichier / Ligne'
    table.rows[1].cells[1].text = f['file']
    colored_cell(table.rows[0].cells[0], '2C3E50')
    table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255,255,255)
    doc.add_paragraph()
    para('Description', bold=True)
    doc.add_paragraph(f['desc'])
    para('Référence CWE', bold=True)
    doc.add_paragraph(f['cwe'])
    para('Scénario d\'exploitation', bold=True)
    doc.add_paragraph(f.get('exploit', 'Voir description.'))
    para('Correctif recommandé', bold=True)
    doc.add_paragraph(f['fix'])
    doc.add_paragraph()

doc.add_page_break()

# ── 5. RECOMMANDATIONS GÉNÉRALES ─────────────────────────────────────
heading('5. Recommandations Générales', 1)

recs = [
    ('Firebase App Check', 'Activer App Check (reCAPTCHA v3) pour bloquer les appels non-autorisés même avec les clés publiques connues.'),
    ('Custom Claims RBAC', 'Migrer tout le contrôle d\'accès vers Firebase Custom Claims. Supprimer toutes les vérifications d\'email hardcodées.'),
    ('Secret Manager', 'Toutes les clés API (Gemini, Meta, LinkedIn) doivent passer par Google Secret Manager. Aucun secret dans Firestore.'),
    ('Content Security Policy', 'Ajouter un en-tête CSP strict via firebase.json pour bloquer les injections XSS et les ressources non-autorisées.'),
    ('Rate Limiting', 'Ajouter un rate limiter sur nexusChat et exchangeSocialToken pour prévenir les abus (Cloud Armor ou compteur Firestore).'),
    ('Rotation des clés', 'Établir une politique de rotation trimestrielle des clés Firebase, FCM VAPID, et clés OAuth.'),
    ('Audit Logs Immuables', 'Exporter les audit_logs vers Cloud Logging ou BigQuery pour garantir la non-répudiation.'),
    ('SAST en CI/CD', 'Intégrer un outil SAST (ex: Semgrep avec ruleset Firebase) dans la pipeline GitHub Actions.'),
    ('Dependency Scanning', 'Activer Dependabot ou npm audit en CI pour détecter automatiquement les CVE dans les dépendances.'),
]

for title, desc in recs:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{title} : ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ── 6. ANNEXE ────────────────────────────────────────────────────────
heading('6. Annexe — Références', 1)
refs = [
    'OWASP Top 10 2021 : https://owasp.org/Top10/',
    'CWE Top 25 2023 : https://cwe.mitre.org/top25/',
    'Firebase Security Best Practices : https://firebase.google.com/docs/rules/basics',
    'Firebase App Check : https://firebase.google.com/docs/app-check',
    'Google Secret Manager : https://cloud.google.com/secret-manager',
]
for r in refs:
    doc.add_paragraph(r, style='List Bullet')

# ── SAVE ─────────────────────────────────────────────────────────────
out = '/Users/yomanraphael/develop/I.P.C/scratch/Audit_Securite_IPC_ERP.docx'
doc.save(out)
print(f'✅ Rapport généré : {out}')
